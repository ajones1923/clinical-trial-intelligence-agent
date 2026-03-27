#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate DOCX files from all Markdown sources in the Clinical Trial
Intelligence Agent docs/ directory.

Applies NVIDIA / HCLS AI Factory branding (VCP palette).  Handles
headings, bold/italic runs, bullet and numbered lists, tables, code
blocks, and horizontal rules.
"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

# -- Paths ----------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
DOCS_DIR = PROJECT_DIR / "docs"

# -- Colors (VCP palette) -------------------------------------
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
HEX_CODE_BG = "F0F0F0"
FONT = "Calibri"
CODE_FONT = "Courier New"

HEADER_TEXT = "HCLS AI Factory \u2014 Clinical Trial Intelligence Agent"
FOOTER_TEXT = "Confidential \u2014 For Internal Use"


# -- Document helpers ------------------------------------------

def new_document():
    """Create a blank Document with standard page setup and base style."""
    doc = Document()

    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

        # Header
        header = s.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(HEADER_TEXT)
        run.font.name = FONT
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_META
        run.font.italic = True

        # Footer with page number
        footer = s.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(FOOTER_TEXT + "    |    Page ")
        run.font.name = FONT
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_META
        # Insert PAGE field
        fld_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
            f'<w:r><w:t>1</w:t></w:r></w:fldSimple>'
        )
        fp._element.append(parse_xml(fld_xml))

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style.font.color.rgb = GRAY_BODY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    return doc


def add_run(p, text, bold=False, italic=False, size=None, color=None, font=None):
    """Add a formatted run to paragraph *p*."""
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def set_paragraph_shading(paragraph, hex_color):
    """Set background shading on a paragraph (used for code blocks)."""
    pPr = paragraph._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{hex_color}"/>')
    pPr.append(shading)


# -- Inline markdown (bold, italic, code, links) ---------------

def add_inline_runs(p, text, base_size=10, base_color=GRAY_BODY):
    """Parse inline markdown (bold, italic, inline code, links) and add runs."""
    # Pattern matches: **bold**, *italic*, `code`, [text](url)
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'   # bold+italic
        r'|(\*\*(.+?)\*\*)'      # bold
        r'|(\*(.+?)\*)'          # italic
        r'|(`(.+?)`)'            # inline code
        r'|(\[([^\]]+)\]\(([^)]+)\))'  # link
    )
    pos = 0
    for m in pattern.finditer(text):
        # Add plain text before this match
        if m.start() > pos:
            add_run(p, text[pos:m.start()], size=base_size, color=base_color)

        if m.group(2):  # bold+italic
            add_run(p, m.group(2), bold=True, italic=True, size=base_size, color=base_color)
        elif m.group(4):  # bold
            add_run(p, m.group(4), bold=True, size=base_size, color=base_color)
        elif m.group(6):  # italic
            add_run(p, m.group(6), italic=True, size=base_size, color=base_color)
        elif m.group(8):  # inline code
            add_run(p, m.group(8), size=9, color=GRAY_CODE, font=CODE_FONT)
        elif m.group(10):  # link
            link_text = m.group(10)
            add_run(p, link_text, size=base_size, color=TEAL)

        pos = m.end()

    # Trailing plain text
    if pos < len(text):
        add_run(p, text[pos:], size=base_size, color=base_color)


# -- Markdown-to-DOCX converter --------------------------------

def convert_md_to_docx(md_path: Path, docx_path: Path):
    """Read a Markdown file and produce a styled DOCX."""
    doc = new_document()
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # -- Code block ----------------------------------------
        if line.strip().startswith("```"):
            lang = line.strip().lstrip("`").strip()
            if lang:
                lp = doc.add_paragraph()
                lp.paragraph_format.space_before = Pt(6)
                lp.paragraph_format.space_after = Pt(0)
                add_run(lp, lang, bold=True, size=8, color=TEAL)
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            # Write code lines
            for cl in code_lines:
                cp = doc.add_paragraph()
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after = Pt(0)
                cp.paragraph_format.line_spacing = 1.0
                cp.paragraph_format.left_indent = Inches(0.2)
                set_paragraph_shading(cp, HEX_CODE_BG)
                add_run(cp, cl if cl else " ", size=9, color=GRAY_CODE, font=CODE_FONT)
            i += 1  # skip closing ```
            continue

        # -- Horizontal rule -----------------------------------
        if re.match(r'^---+\s*$', line.strip()):
            # Add a thin horizontal line via a bottom-bordered paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:color="{HEX_TEAL}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # -- Table ---------------------------------------------
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                _add_table_from_md(doc, table_lines)
            continue

        # -- Title (# ) ----------------------------------------
        if line.startswith("# ") and not line.startswith("## "):
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, title_text, bold=True, size=24, color=NAVY)
            i += 1
            continue

        # -- Heading 1 (## ) -----------------------------------
        if line.startswith("## "):
            h_text = line[3:].strip()
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(8)
            for r in p.runs:
                r.clear()
            add_run(p, h_text, bold=True, size=16, color=NAVY)
            i += 1
            continue

        # -- Heading 2 (### ) ----------------------------------
        if line.startswith("### "):
            h_text = line[4:].strip()
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                r.clear()
            add_run(p, h_text, bold=True, size=13, color=TEAL)
            i += 1
            continue

        # -- Heading 3 (#### ) ---------------------------------
        if line.startswith("#### "):
            h_text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            add_run(p, h_text, bold=True, size=11, color=GREEN)
            i += 1
            continue

        # -- Numbered list -------------------------------------
        m_num = re.match(r'^(\d+)\.\s+(.*)', line)
        if m_num:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            add_run(p, f"{m_num.group(1)}.  ", bold=True, size=10, color=GRAY_BODY)
            add_inline_runs(p, m_num.group(2), base_size=10, base_color=GRAY_BODY)
            i += 1
            continue

        # -- Bullet list ---------------------------------------
        m_bullet = re.match(r'^[-*]\s+(.*)', line)
        if m_bullet:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            add_run(p, "\u2022  ", size=10, color=TEAL)
            add_inline_runs(p, m_bullet.group(1), base_size=10, base_color=GRAY_BODY)
            i += 1
            continue

        # -- Indented bullet (sub-bullet) ----------------------
        m_sub = re.match(r'^  [-*]\s+(.*)', line)
        if m_sub:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.6)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            add_run(p, "\u25E6  ", size=9, color=TEAL)
            add_inline_runs(p, m_sub.group(1), base_size=10, base_color=GRAY_BODY)
            i += 1
            continue

        # -- Blank line ----------------------------------------
        if not line.strip():
            i += 1
            continue

        # -- Body paragraph ------------------------------------
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        add_inline_runs(p, line, base_size=10, base_color=GRAY_BODY)
        i += 1

    doc.save(str(docx_path))


def _add_table_from_md(doc, table_lines):
    """Parse markdown table lines and create a styled DOCX table."""
    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        return cells

    # First line = headers
    headers = parse_row(table_lines[0])

    # Second line = separator (skip it)
    # Remaining lines = data rows
    rows = []
    for tl in table_lines[2:]:
        row = parse_row(tl)
        # Pad or trim to match header count
        while len(row) < len(headers):
            row.append("")
        rows.append(row[:len(headers)])

    if not headers:
        return

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True

    # Header row
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=9, color=WHITE)

    # Data rows
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                add_run(p, val, bold=True, size=9, color=TEAL)
            else:
                add_run(p, val, size=9, color=GRAY_BODY)

    # Spacer after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(4)


# -- Main ------------------------------------------------------

def main():
    print(f"Docs directory: {DOCS_DIR}")

    # Discover all .md files in the docs directory
    md_files = sorted(DOCS_DIR.glob("*.md"))

    if not md_files:
        print("No .md files found in docs/ directory.")
        return 1

    print(f"Processing {len(md_files)} markdown file(s)...\n")

    success = 0
    errors = []

    for md_path in md_files:
        md_name = md_path.name
        docx_name = md_name.replace(".md", ".docx")
        docx_path = DOCS_DIR / docx_name

        try:
            convert_md_to_docx(md_path, docx_path)
            size_kb = docx_path.stat().st_size / 1024
            print(f"  OK    {docx_name}  ({size_kb:.0f} KB)")
            success += 1
        except Exception as exc:
            msg = f"  FAIL  {md_name}: {exc}"
            print(msg)
            errors.append(msg)

    print(f"\nDone: {success}/{len(md_files)} files converted.")
    if errors:
        print("Errors:")
        for e in errors:
            print(f"  {e}")
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
